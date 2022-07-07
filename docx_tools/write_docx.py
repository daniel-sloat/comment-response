import docx
from docx.enum.text import WD_UNDERLINE
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt

from logtools import logtools


def _create_styles(doc: docx.Document) -> None:
    def _style_maker(
        doc: docx.Document,
        name: str,
        base_style: str = "Normal",
        left_indent: int | float = 0,
        space_before: int = 12,
        space_after: int = 12,
        next_style: str = "",
        keep_with_next: bool = False,
    ) -> None:
        styles = doc.styles
        style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles[base_style]
        style.paragraph_format.left_indent = Inches(left_indent)
        style.paragraph_format.space_before = Pt(space_before)
        style.paragraph_format.space_after = Pt(space_after)
        if next_style:
            style.next_paragraph_style = styles[next_style]
        if keep_with_next:
            style.paragraph_format.keep_with_next = keep_with_next
        return None

    _style_maker(doc, "Comments")
    _style_maker(doc, "Response", left_indent=0.5, next_style="Response")


def _word_formats(tag: str, run) -> None:
    for f in tag:
        match f:
            case "b":
                run.font.bold = True
            case "i":
                run.font.italic = True
            case "u":
                run.font.underline = True
            case "w":
                run.font.underline = WD_UNDERLINE.DOUBLE
            case "s":
                run.font.strike = True
            case "z":
                run.font.double_strike = True
            case "x":
                run.font.superscript = True
            case "v":
                run.font.subscript = True
    return None


def _write_comments_and_responses(
    doc,
    group_data,
    comment_intro_every_comment=True,
    comment_intro="Comment",
    response_intro="Response",
    intro_sep=": ",
):
    # COMMENTS
    for comment in (comments := group_data["comment_data"]["comments"]):
        paragraph = doc.add_paragraph(style="Comments")
        if comment_intro_every_comment or len(comments) > 1:
            intro = paragraph.add_run(comment_intro)
            intro.underline = True
            paragraph.add_run(intro_sep)
        for para_no, para in enumerate(comment.paragraphs):
            if para_no == 0:
                for run in para.runs:
                    r = paragraph.add_run(run.text)
                    _word_formats(run.props, r)
            else:
                paragraph = doc.add_paragraph(style="Comments")
                for run in para.runs:
                    r = paragraph.add_run(run.text)
                    _word_formats(run.props, r)
    # RESPONSE
    for response in group_data["comment_data"]["response"]:
        paragraph = doc.add_paragraph(style="Response")
        intro = paragraph.add_run(response_intro)
        intro.italic = True
        intro.bold = True
        paragraph.add_run(intro_sep)
        for para_no, para in enumerate(response.paragraphs):
            if para_no == 0:
                for run in para.runs:
                    r = paragraph.add_run(run.text)
                    _word_formats(run.props, r)
            else:
                paragraph = doc.add_paragraph(style="Response")
                for run in para.runs:
                    r = paragraph.add_run(run.text)
                    _word_formats(run.props, r)


def multiple_comments(
    group,
    indicate_quantity,
    single_comment,
    multiple,
):
    if indicate_quantity:
        if len(group["data"]["comment_data"]["comments"]) > 1:
            return multiple
        else:
            return single_comment
    else:
        return ""


def _write_document(
    doc: docx.Document,
    top_level: list,
    outline_level_start: int,
    indicate_quantity=True,
    single_comment="Comment: ",
    multiple="Multiple Comments: ",
    **kwargs,
) -> None:
    def recursion(top_level, outline_level=1, indicate_quantity=True):
        for group in top_level:

            if isinstance(group.get("data"), dict):
                quantity = multiple_comments(group, indicate_quantity, single_comment, multiple)
                doc.add_heading(quantity + group["heading"], outline_level)
                _write_comments_and_responses(doc, group["data"], **kwargs)
            else:
                doc.add_heading(group["heading"], outline_level)
                recursion(group.get("data"), outline_level + 1, indicate_quantity)

    recursion(top_level, outline_level_start, indicate_quantity)
    return None


@logtools.log_write_docx
def commentsectiondoc(
    nested_comment_responses: list,
    savename: str = "output\CommentResponse.docx",
    **kwargs,
) -> None:
    doc = docx.Document()
    _create_styles(doc)
    _write_document(doc, nested_comment_responses, **kwargs)
    doc.save(savename)
    return savename

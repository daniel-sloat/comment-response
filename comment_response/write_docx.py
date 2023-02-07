"""Write comment-response section."""

from reprlib import Repr

import docx
from docx.document import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_UNDERLINE
from docx.shared import Inches, Pt
from comment_response.group.group_records import GroupRecords

from comment_response.records import Records
from xlsx.sheets.datasheet import DataSheet

PARAGRAPH_STYLE = WD_STYLE_TYPE.PARAGRAPH  # pylint: disable=no-member
DOUBLE_UNDERLINE_STYLE = WD_UNDERLINE.DOUBLE  # pylint: disable=no-member


def style_maker(
    doc: docx.Document,
    name: str,
    base_style: str = "Normal",
    left_indent: int | float = 0,
    space_before: int = 12,
    space_after: int = 12,
    next_style: str = "",
    keep_with_next: bool = False,
) -> None:
    styles = doc.styles
    style = styles.add_style(name, PARAGRAPH_STYLE)
    style.base_style = styles[base_style]
    style.paragraph_format.left_indent = Inches(left_indent)
    style.paragraph_format.space_before = Pt(space_before)
    style.paragraph_format.space_after = Pt(space_after)
    if next_style:
        style.next_paragraph_style = styles[next_style]
    if keep_with_next:
        style.paragraph_format.keep_with_next = keep_with_next
    return None


def word_formats(tag: dict | None, run) -> None:
    if tag:
        match tag:
            case {"rPr": {"b": value_dict}}:
                run.font.bold = True
            case {"rPr": {"i": value_dict}}:
                run.font.italic = True
            case {"rPr": {"u": value_dict}}:
                run.font.underline = True
                # run.font.underline = DOUBLE_UNDERLINE_STYLE
            case {"rPr": {"strike": value_dict}}:
                run.font.strike = True
            case {"rPr": {"strike": value_dict, "color": {"rgb": "FFFF0000"}}}:
                run.font.double_strike = True
            case {"rPr": {"vertAlign": value_dict}}:
                run.font.superscript = True
                # run.font.subscript = True


def write_comments(document, records):
    for comment in records.comments:
        paragraph = document.add_paragraph(style="Comments")
        if len(records.comments) > 1:
            intro = paragraph.add_run("Comment")
            intro.underline = True
            paragraph.add_run(": ")
        for para_no, para in enumerate(comment):
            if para_no == 0:
                for run in para:
                    text, props = run
                    added_run = paragraph.add_run(text)
                    word_formats(props, added_run)
            else:
                paragraph = document.add_paragraph(style="Comments")
                for run in para:
                    text, props = run
                    added_run = paragraph.add_run(text)
                    word_formats(props, added_run)


def write_response(document, records):
    paragraph = document.add_paragraph(style="Response")
    intro = paragraph.add_run("Agency Response")
    intro.italic = True
    intro.bold = True
    paragraph.add_run(": ")
    for para_no, para in enumerate(records.response):
        if para_no == 0:
            for run in para:
                text, props = run
                added_run = paragraph.add_run(text)
                word_formats(props, added_run)
        else:
            paragraph = document.add_paragraph(style="Response")
            for run in para:
                text, props = run
                added_run = paragraph.add_run(text)
                word_formats(props, added_run)


def recursive_write(document: Document, grouped_records, config, outline_level=0):
    outline_level += 1
    for item in grouped_records:
        match item:
            case {"heading": heading, "data": [{"records": records}]}:
                # Base case (normal)
                multiple = len(records) > 1
                if multiple:
                    document.add_heading(
                        f"Multiple Comments: {heading.title}", level=outline_level
                    )
                else:
                    document.add_heading(
                        f"Comment: {heading.title}", level=outline_level
                    )

                records = Records(records, config)
                write_comments(document, records)
                write_response(document, records)

            case {"heading": heading, "data": data}:
                # Recursive case (only writes heading)
                document.add_heading(heading.title, level=outline_level)
                recursive_write(document, data, config, outline_level)

            case {"records": records}:
                # Base case (for when records are not fully classified)
                records = Records(records, config)
                write_comments(document, records)
                write_response(document, records)


cs_repr = Repr()
cs_repr.maxlevel = 2
cs_repr.maxdict = 5
cs_repr.maxlist = 1


class CommentSection:
    """CommentSection object for writing comment-response section to docx."""

    def __init__(self, sheet, **config):
        self._sheet: DataSheet = sheet
        self.config = config
        self.outline_level = 0
        self.group_records = GroupRecords(self._sheet, **config)

    def __repr__(self):
        return f"{self.__class__.__name__}(sheetname={self._sheet._name})"

    def write(self, filename="output/commentsection.docx"):
        doc: Document = docx.Document()
        style_maker(doc, "Comments")
        style_maker(doc, "Response", left_indent=0.5, next_style="Response")
        recursive_write(
            doc, self.group_records.group(), self.config, self.outline_level
        )
        doc.save(filename)

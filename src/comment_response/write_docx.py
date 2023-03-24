"""Write comment-response section."""

from pathlib import Path

import docx
from docx.document import Document
from docx.enum.base import EnumValue
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_UNDERLINE
from docx.shared import Inches, Pt
from docx.styles.style import _ParagraphStyle
from docx.styles.styles import Styles
from docx.text.run import Run

from comment_response.group.colsort import ColSort
from comment_response.group.group_records import GroupRecords
from comment_response.records import Records

# from xlsx.sheets.datasheet import DataSheet

PARAGRAPH_STYLE: EnumValue = WD_STYLE_TYPE.PARAGRAPH  # pylint: disable=no-member
DOUBLE_UNDERLINE_STYLE: EnumValue = WD_UNDERLINE.DOUBLE  # pylint: disable=no-member


def style_maker(
    doc: Document,
    name: str,
    base_style: str = "Normal",
    left_indent: int | float = 0,
    space_before: int = 12,
    space_after: int = 12,
    next_style: str = "",
    keep_with_next: bool = False,
) -> None:
    styles: Styles = doc.styles
    style: _ParagraphStyle = styles.add_style(name, PARAGRAPH_STYLE)
    style.base_style = styles[base_style]
    style.paragraph_format.left_indent = Inches(left_indent)
    style.paragraph_format.space_before = Pt(space_before)
    style.paragraph_format.space_after = Pt(space_after)
    if next_style:
        style.next_paragraph_style = styles[next_style]
    if keep_with_next:
        style.paragraph_format.keep_with_next = keep_with_next


def word_formats(tag: dict | None, run: Run) -> None:
    def toggle(value_dict: dict[str:str]) -> bool:
        match value_dict:
            case {}:
                return True
            case {"val": value}:
                try:
                    if str(value).casefold() == "true":
                        return True
                    return False
                except TypeError:
                    return bool(int(value))
            case _:
                return False

    if "b" in tag:
        run.font.bold = toggle(tag["b"])
    if "i" in tag:
        run.font.italic = toggle(tag["i"])
    if "u" in tag:
        match tag["u"]:
            case {"val": _type}:
                string = str(_type).casefold()
                if string == "double" or string == "wavydouble":
                    run.font.underline = DOUBLE_UNDERLINE_STYLE
                elif string == "none":
                    run.font.underline = False
                else:
                    run.font.underline = True
            case _:
                run.font.underline = toggle(tag["u"])
    if "strike" in tag:
        if tag.get("color", {}).get("rgb") == "FFFF0000":
            run.font.double_strike = toggle(tag["strike"])
        else:
            run.font.strike = toggle(tag["strike"])
    if "vertAlign" in tag:
        match tag["vertAlign"]:
            case {"val": "superscript"}:
                run.font.superscript = True
            case {"val": "subscript"}:
                run.font.subscript = True


def write_comments(document: Document, records: Records, config: dict) -> None:
    for comment in records.comments:
        paragraph = document.add_paragraph(style="Comments")
        if (
            len(records.comments) > 1
            or config["doc"]["custom"]["comment_intro_every_comment"]
        ):
            intro = paragraph.add_run(config["doc"]["custom"]["comment_intro"])
            intro.underline = True
            paragraph.add_run(config["doc"]["custom"]["intro_sep"])
        for para_no, para in enumerate(comment):
            if para_no == 0:
                for run in para:
                    added_run = paragraph.add_run(run.text)
                    if run.props:
                        word_formats(run.props, added_run)
            else:
                paragraph = document.add_paragraph(style="Comments")
                for run in para:
                    added_run = paragraph.add_run(run.text)
                    if run.props:
                        word_formats(run.props, added_run)


def write_response(document: Document, records: Records, config: dict) -> None:
    paragraph = document.add_paragraph(style="Response")
    intro = paragraph.add_run(config["doc"]["custom"]["response_intro"])
    intro.italic = True
    intro.bold = True
    paragraph.add_run(config["doc"]["custom"]["intro_sep"])
    for para_no, para in enumerate(records.response):
        if para_no == 0:
            for run in para:
                added_run = paragraph.add_run(run.text)
                if run.props:
                    word_formats(run.props, added_run)
        else:
            paragraph = document.add_paragraph(style="Response")
            for run in para:
                added_run = paragraph.add_run(run.text)
                if run.props:
                    word_formats(run.props, added_run)


def indicate_quantity(records: Records, config: dict) -> str:
    if config["doc"]["custom"]["indicate_quantity"]:
        multiple = len(records.comments) > 1
        if multiple:
            return config["doc"]["custom"]["multiple_comments"]
        return config["doc"]["custom"]["single_comment"]
    return ""


def recursive_write(
    document: Document,
    grouped_records: GroupRecords,
    config: dict,
    outline_level: int = 0,
):
    outline_level += 1
    for item in grouped_records:
        match item:
            case {"heading": ColSort() as heading, "data": [{"records": records}]}:
                # Base case (normal)
                records = Records(records, config)
                pre = indicate_quantity(records, config)
                document.add_heading(f"{pre}{heading.title}", level=outline_level)
                write_comments(document, records, config)
                write_response(document, records, config)

            case {"records": records}:
                # Base case (for when records are not fully classified)
                records = Records(records, config)
                write_comments(document, records, config)
                write_response(document, records, config)

            case {"heading": ColSort() as heading, "data": data}:
                # Recursive case (only writes heading)
                document.add_heading(heading.title, level=outline_level)
                recursive_write(document, data, config, outline_level)


class CommentSection:
    """CommentSection object for writing comment-response section to docx."""

    def __init__(self, sheet, **config):
        self._sheet = sheet
        self.config: dict = config
        self.outline_level: int = self.config["doc"]["custom"]["outline_level_start"]
        self.group_records = GroupRecords(self._sheet, **config)
        self.filename = Path(self.config["doc"]["savename"])

    def __repr__(self):
        return f"{self.__class__.__name__}(sheetname={self._sheet._name})"

    def write(self):
        doc: Document = docx.Document()
        style_maker(doc, "Comments")
        style_maker(doc, "Response", left_indent=0.5, next_style="Response")
        recursive_write(
            doc, self.group_records.group(), self.config, self.outline_level
        )
        self.filename.parent.mkdir(exist_ok=True)
        doc.save(self.filename)

"""Write comment-response section."""

from docx.document import Document

from comment_response.group.recursive_group import Heading
from comment_response.parts.comment_group import CommentGroup
from comment_response.write.format_adapter import format_adapter


def write_comments(
    document: Document, records: CommentGroup, custom_config: dict
) -> None:
    for comment in records.comments:
        paragraph = document.add_paragraph(style="Comments")
        if len(records.comments) > 1 or custom_config["comment_intro_every_comment"]:
            intro = paragraph.add_run(custom_config["comment_intro"])
            intro.underline = True
            paragraph.add_run(custom_config["intro_sep"])
        for para in (paras := comment.paragraphs):
            if para == paras[0]:
                for run in para.runs:
                    added_run = paragraph.add_run(run.text)
                    if run.props:
                        format_adapter(run.props, added_run)
            else:
                paragraph = document.add_paragraph(style="Comments")
                for run in para.runs:
                    added_run = paragraph.add_run(run.text)
                    if run.props:
                        format_adapter(run.props, added_run)
        paragraph.add_run(f" ({comment.tag})")


def write_response(
    document: Document, records: CommentGroup, custom_config: dict
) -> None:
    paragraph = document.add_paragraph(style="Response")
    intro = paragraph.add_run(custom_config["response_intro"])
    intro.italic = True
    intro.bold = True
    paragraph.add_run(custom_config["intro_sep"])
    for para_no, para in enumerate(records.response.paragraphs):
        if para_no == 0:
            for run in para.runs:
                added_run = paragraph.add_run(run.text)
                if run.props:
                    format_adapter(run.props, added_run)
        else:
            paragraph = document.add_paragraph(style="Response")
            for run in para.runs:
                added_run = paragraph.add_run(run.text)
                if run.props:
                    format_adapter(run.props, added_run)


def indicate_quantity(records: CommentGroup, quantity_config: dict) -> str:
    if quantity_config["indicate_quantity"]:
        multiple = len(records.comments) > 1
        if multiple:
            return quantity_config["multiple_comments"]
        return quantity_config["single_comment"]
    return ""


def recursive_write(
    document: Document,
    grouped_records: list[dict],
    config: dict,
    outline_level: int = 0,
):
    """Recursively write comments and response section."""
    outline_level += 1
    for item in grouped_records:
        match item:
            case {"heading": Heading() as heading, "data": [{"records": records}]}:
                # Base case (normal)
                records = CommentGroup(records, config)
                pre = indicate_quantity(records, config["other"]["quantity"])
                document.add_heading(f"{pre}{heading.title}", level=outline_level)
                write_comments(document, records, config["other"]["custom"])
                write_response(document, records, config["other"]["custom"])

            case {"records": records}:
                # Base case (for when records are not fully classified)
                records = CommentGroup(records, config)
                write_comments(document, records, config["other"]["custom"])
                write_response(document, records, config["other"]["custom"])

            case {"heading": Heading() as heading, "data": data}:
                # Recursive case (only writes heading)
                document.add_heading(heading.title, level=outline_level)
                recursive_write(document, data, config, outline_level)

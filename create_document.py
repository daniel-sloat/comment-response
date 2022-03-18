#!/usr/bin/env python3.10
# -*- coding: utf-8 -*-

import docx
from docx.enum.text import WD_UNDERLINE
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt

def _create_styles(doc: docx.Document) -> None:
    def _style_maker(
        doc: docx.Document,
        name: str,
        base_style: str="Normal",
        left_indent: int | float=0,
        space_before: int=12,
        space_after: int=12,
        next_style: str="",
        keep_with_next: bool=False
    ) -> None:
        styles = doc.styles
        style = styles.add_style(name,WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles[base_style]
        style.paragraph_format.left_indent = Inches(left_indent)
        style.paragraph_format.space_before = Pt(space_before)
        style.paragraph_format.space_after = Pt(space_after)
        if next_style:
            style.next_paragraph_style = styles[next_style]
        if keep_with_next:
            style.paragraph_format.keep_with_next = keep_with_next
        return None

    _style_maker(doc,"Comments")
    _style_maker(doc,"Response",left_indent=0.5,next_style="Response")
    _style_maker(doc,"AgencyResponse",
                left_indent=0.5,space_after=2,
                next_style="Response",keep_with_next=True)

def _write_comments(
    doc: docx.Document,
    nested_comment_responses: list
) -> None:
    for section1_data, section1_name in nested_comment_responses:
            doc.add_heading(section1_name, 1)
            for section2_data, section2_name in section1_data:
                doc.add_heading(section2_name, 2)
                for section3_data, section3_name, section3_response in section2_data:
                    if len(section3_data) > 1:
                        plural_comments = "Multiple Comments:"
                    else:
                        plural_comments = "Comment:"
                    doc.add_heading(f"{plural_comments} {section3_name}", 3)
                    for comment in section3_data:
                        paragraph = doc.add_paragraph()
                        paragraph.style = "Comments"
                        if len(section3_data) > 1:
                            paragraph.add_run("Comment", style="Run u")
                            paragraph.add_run(": ", style="Run ")
                        for para_no, para in enumerate(comment):
                            if para_no == 0:
                                for run in para:
                                    paragraph.add_run(run[1], style="Run " + run[0])                                                          
                            else:
                                paragraph = doc.add_paragraph()
                                paragraph.style = "Comments"
                                for run in para:
                                    paragraph.add_run(run[1], style="Run " + run[0])
                    paragraph = doc.add_paragraph()
                    paragraph.style = "AgencyResponse"
                    paragraph.add_run("Agency Response", style="Run biu")
                    paragraph.add_run(": ", style="Run ")
                    # Rich text response. There is only one response, so one less 
                    # level of iteration than comments.
                    for para_no, para in enumerate(section3_response):
                        #paragraph = doc.add_paragraph()
                        paragraph.style = "Response"
                        if para_no == 0:
                            for run in para:
                                paragraph.add_run(run[1], style="Run " + run[0])                                                          
                        else:
                            paragraph = doc.add_paragraph()
                            paragraph.style = "Response"
                            for run in para:
                                paragraph.add_run(run[1], style="Run " + run[0])
    return None
          
def _word_formats(
    doc: docx.Document,
    formats: list[str],
    add_styles: list[str]
) -> None:
    styles = doc.styles
    if add_styles:
        for style in add_styles:
            formats.append(style)
        formats = list(set(formats))
    for tag in formats:
        charstyle_font = styles.add_style("Run " + tag, WD_STYLE_TYPE.CHARACTER).font
        for v in tag:
            if v == "b": charstyle_font.bold = True
            elif v == "i": charstyle_font.italic = True
            elif v == "u": charstyle_font.underline = True
            elif v == "w": charstyle_font.underline = WD_UNDERLINE.DOUBLE
            elif v == "s": charstyle_font.strike = True
            elif v == "z": charstyle_font.double_strike = True
            elif v == "x": charstyle_font.superscript = True
            elif v == "v": charstyle_font.subscript = True
    return None          

def commentsectiondoc(
    nested_comment_responses: list,
    formats: list[str],
    savename: str="CommentResponseSection.docx"
) -> None:
    print("Creating Comments and Response section document... ")
    doc = docx.Document()
    _create_styles(doc)
    # Need to include default syles "u" and "biu" 
    # because they are used to write text in word doc
    _word_formats(doc,formats,add_styles=["u","biu"])
    _write_comments(doc,nested_comment_responses)
    doc.save(savename)
    print("Comments and response section document created: " + savename)
    return None

def automarkdoc(
    entry_list: list,
    savename: str="AutoMark.docx"
) -> None:
    print("Creating AutoMark document... ")
    def _write_table(doc,entry_list):
        # AutoMark document is document with two col table for automatically
        # marking index entries in another document.
        # Use table._cells to "pop" out the cells from the table, limiting 
        # the amount of calls to the table in the Word document (improving 
        # speed by multiple times). Updates Word document only after the 
        # table is filled.
        # https://theprogrammingexpert.com/write-table-fast-python-docx/
        table = doc.add_table(rows=len(entry_list), cols=2)
        table_cells = table._cells
        for i in range(len(entry_list)):
            for j in range(len(entry_list[i])):
                table_cells[j + i * 2].text = str(entry_list[i][j])

    doc = docx.Document()
    _write_table(doc,entry_list)
    doc.save(savename)
    print("AutoMark document created: " + savename)
    return None